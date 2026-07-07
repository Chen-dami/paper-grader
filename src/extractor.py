"""
Word 文档解析器 —— 把 .docx 拆成结构化数据。
输出：文字段落、表格内容、图片路径、嵌入文件路径。
"""
import os
import json
import zipfile
import shutil
from docx import Document


def extract(docx_path: str, output_dir: str = "output") -> dict:
    """
    解析一份 docx 试卷，返回结构化字典。

    返回结构：
    {
        "file_name": "xxx.docx",
        "student_info": {"学号": "", "姓名": "", "班级": ""},
        "paragraphs": [(index, text), ...],
        "tables": [[{cell_text}, ...], ...],   # 每个表格是一行行数据
        "images": [(local_path, width, height, size_bytes), ...],
        "embedded_files": [(local_path, ext), ...],
    }
    """
    file_name = os.path.basename(docx_path)
    paper_dir = os.path.join(output_dir, os.path.splitext(file_name)[0])
    img_dir = os.path.join(paper_dir, "images")
    embed_dir = os.path.join(paper_dir, "embeddings")
    os.makedirs(img_dir, exist_ok=True)
    os.makedirs(embed_dir, exist_ok=True)

    doc = Document(docx_path)

    # ---- 1. 提取文字段落 ----
    paragraphs = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        paragraphs.append((i, text))

    # ---- 2. 提取表格 ----
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    # ---- 3. 提取图片（按文档中的出现顺序，并记录所在表位置） ----
    images = []
    # 获取图片在文档中出现的顺序，及其附近表格索引
    # img_order: {media_filename: (order_idx, table_index)}
    img_order = _get_image_document_order(docx_path)
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if 'media' in name and not name.endswith('/'):
                data = z.read(name)
                fname = os.path.basename(name)
                local_path = os.path.join(img_dir, fname)

                # 跳过重复提取
                if not os.path.exists(local_path):
                    with open(local_path, 'wb') as f:
                        f.write(data)

                # 读取图片尺寸
                width, height = _get_image_size(data, fname)
                # order_idx: 图片在文档中的序号, table_idx: 最近的表格索引
                order_idx, table_idx = img_order.get(fname, (-1, -1))
                images.append((local_path, width, height, len(data), order_idx, table_idx))

    # 按 order_idx 排序，确保图片按文档顺序排列
    images.sort(key=lambda x: x[4] if len(x) >= 5 and x[4] >= 0 else 99999)

    # ---- 4. 提取嵌入文件（OLE 对象） ----
    embedded_files = []
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if 'embedding' in name.lower() and not name.endswith('/'):
                data = z.read(name)
                fname = os.path.basename(name)
                local_path = os.path.join(embed_dir, fname)

                if not os.path.exists(local_path):
                    with open(local_path, 'wb') as f:
                        f.write(data)

                # 尝试提取嵌套的 Office 文件
                nested = _extract_ole(data, embed_dir, fname)
                if nested:
                    embedded_files.append((nested, os.path.splitext(nested)[1]))

    # ---- 5. 解析学生信息 ----
    student_info = _parse_student_info(paragraphs, tables, docx_path)

    return {
        "file_name": file_name,
        "paper_dir": paper_dir,
        "student_info": student_info,
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
        "embedded_files": embedded_files,
    }


def _get_image_size(data: bytes, fname: str) -> tuple:
    """从 PNG/JPEG 二进制数据中读取宽高"""
    import struct
    try:
        if fname.lower().endswith('.png'):
            w = struct.unpack('>I', data[16:20])[0]
            h = struct.unpack('>I', data[20:24])[0]
            return w, h
        elif fname.lower().endswith(('.jpg', '.jpeg')):
            # JPEG 尺寸解析较复杂，简化处理
            return 0, 0
        else:
            return 0, 0
    except Exception:
        return 0, 0


def _extract_ole(data: bytes, output_dir: str, base_name: str) -> str | None:
    """
    从 OLE 对象中提取嵌套文件（Office 文档或视频）。
    优先从 Ole10Native 流提取原始数据，再检测类型。
    """
    # 1. 先尝试提取 Ole10Native 原始文件数据（真正的嵌入文件在这里）
    native = _try_extract_ole_native(data)

    # 2. 如果取到了 native，检测类型
    if native and len(native) > 100:
        # 视频检测（ftyp 可能在 OLE 包装头之后，搜索前 500 字节）
        head500 = native[:500]
        if b'ftyp' in head500:
            ext = '.mp4'
        elif head500[:4] == b'RIFF':
            ext = '.avi'
        else:
            # 尝试从原始 OLE 数据中找 ZIP（Office 文件）
            ext = _detect_ole_ext(data)

        out_path = os.path.join(output_dir, base_name.replace('.bin', ext))
        if not os.path.exists(out_path):
            with open(out_path, 'wb') as f:
                f.write(native)
        return out_path

    # 3. 没有 Ole10Native，从原始数据扫描
    for sig, ext in [(b'ftyp', '.mp4'), (b'RIFF', '.avi'), (b'moov', '.mov')]:
        idx = data.find(sig, 4, 512)
        if idx > 0:
            out_path = os.path.join(output_dir, base_name.replace('.bin', ext))
            if not os.path.exists(out_path):
                with open(out_path, 'wb') as f:
                    f.write(data)
            return out_path

    # 4. ZIP 头 → Office 文件
    idx = data.find(b'PK\x03\x04')
    if idx >= 4:
        zip_data = data[idx:]
        if zip_data[:4] == b'PK\x03\x04':
            ext = _detect_office_ext(zip_data)
            out_path = os.path.join(output_dir, base_name.replace('.bin', ext))
            if not os.path.exists(out_path):
                with open(out_path, 'wb') as f:
                    f.write(zip_data)
            return out_path

    return None


def _detect_ole_ext(data: bytes) -> str:
    """从 OLE 原始数据中检测嵌入文件类型"""
    for sig, ext in [(b'ftyp', '.mp4'), (b'RIFF', '.avi'), (b'moov', '.mov')]:
        if sig in data[4:512]:
            return ext
    # 检查 ZIP
    idx = data.find(b'PK\x03\x04')
    if idx >= 4:
        return _detect_office_ext(data[idx:])
    return '.bin'


def _detect_office_ext(zip_data: bytes) -> str:
    """从 ZIP 内容判断是哪种 Office 文件"""
    head = zip_data[:5000]
    if b'xl/' in head:
        return '.xlsx'
    elif b'word/' in head:
        return '.docx'
    elif b'ppt/' in head:
        return '.pptx'
    return '.zip'


def _try_extract_ole_native(data: bytes) -> bytes | None:
    """
    从 OLE2 复合文档的 Ole10Native 流中提取原始嵌入文件。
    使用 olefile 库解析 OLE2 结构。
    """
    if data[:8] != b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
        return None
    try:
        import olefile
        ole = olefile.OleFileIO(data)
        # 查找 Ole10Native 流
        for s in ole.listdir():
            if 'Ole10Native' in '/'.join(s):
                raw = ole.openstream(s).read()
                ole.close()
                # Ole10Native 结构：4字节大小 + 嵌入文件数据
                if len(raw) > 8:
                    import struct
                    size = struct.unpack_from('<I', raw, 0)[0]
                    if 100 < size < len(raw):
                        return raw[4:4 + size]
                    # 有时候原生数据直接从 offset 4 开始
                    return raw[4:]
                return raw
        ole.close()
    except Exception:
        pass
    return None


def extract_from_student_folder(docx_path: str, output_dir: str = "output",
                                supplementary_files: dict | None = None) -> dict:
    """
    从学生文件夹提取试卷 —— 正常解析 docx，同时将文件夹中的
    辅助文件（图片、Excel）注入到提取结果中。

    supplementary_files = {
        "images": ["/path/to/img1.png", ...],
        "excel": ["/path/to/data.xlsx", ...],
    }
    """
    paper = extract(docx_path, output_dir)
    paper_dir = paper["paper_dir"]

    if supplementary_files:
        img_dir = os.path.join(paper_dir, "images")
        embed_dir = os.path.join(paper_dir, "embeddings")
        os.makedirs(img_dir, exist_ok=True)
        os.makedirs(embed_dir, exist_ok=True)

        # 文件夹素材插入到列表开头，优先于 docx 内嵌文件
        for img_path in supplementary_files.get("images", []):
            if not os.path.exists(img_path):
                continue
            fname = os.path.basename(img_path)
            dest = os.path.join(img_dir, fname)
            if not os.path.exists(dest):
                shutil.copy2(img_path, dest)
            size = os.path.getsize(dest)
            w, h = _get_image_size(open(dest, "rb").read(), fname)
            paper["images"].insert(0, (dest, w, h, size, -1, -1))  # -1 = 位置未知（外部文件）

        for xlsx_path in supplementary_files.get("excel", []):
            if not os.path.exists(xlsx_path):
                continue
            fname = os.path.basename(xlsx_path)
            dest = os.path.join(embed_dir, fname)
            if not os.path.exists(dest):
                shutil.copy2(xlsx_path, dest)
            paper["embedded_files"].insert(0, (dest, os.path.splitext(fname)[1]))

    return paper


def _parse_student_info(paragraphs: list, tables: list, docx_path: str = "") -> dict:
    """从段落、表格、文件夹名中提取学生信息"""
    import re
    info = {"学号": "", "姓名": "", "班级": ""}

    # 从段落中查找
    for i, text in paragraphs:
        if '班级' in text and '姓名' in text and '学号' in text:
            for field in ['班级', '姓名', '学号']:
                pattern = rf'{field}\s*(\S+)'
                m = re.search(pattern, text)
                if m:
                    info[field] = m.group(1).strip()
            break

    # 从表格中查找
    if not info.get("学号"):
        for table in tables:
            for row in table:
                row_text = ' '.join(row)
                if '学号' in row_text or '姓名' in row_text:
                    for i, cell in enumerate(row):
                        if '学号' in cell and i + 1 < len(row):
                            info['学号'] = row[i + 1].strip()
                        if '姓名' in cell and i + 1 < len(row):
                            info['姓名'] = row[i + 1].strip()
                        if '班级' in cell and i + 1 < len(row):
                            info['班级'] = row[i + 1].strip()

    # 兜底：从文件夹名提取（如 255102030101李杰鸿 或 255102030101_李杰鸿）
    if not info.get("学号") and docx_path:
        import os
        parent = os.path.basename(os.path.dirname(docx_path))
        sid_match = re.search(r'(\d{11,12})', parent)
        if sid_match:
            info['学号'] = sid_match.group(1)
        name_match = re.search(r'[一-鿿]{2,4}', parent)
        if name_match:
            info['姓名'] = name_match.group(0)
        # 班级：从上级目录名提取（如 data/papers/软件2501/学生文件夹/docx）
        grandparent = os.path.basename(os.path.dirname(os.path.dirname(docx_path)))
        if re.match(r'^[一-鿿_a-zA-Z]+\d{2,}', grandparent):
            info['班级'] = grandparent

    return info


def _get_image_document_order(docx_path: str) -> dict:
    """
    解析 document.xml 获取图片在文档中的出现顺序及所在表格。
    返回: {media_filename: (order_idx, table_idx)}
    - order_idx: 图片在文档中出现的序号（0-based，越小越靠前）
    - table_idx: 图片所在的表格索引（对应 doc.tables 的索引），-1 表示不在表格内
    """
    import zipfile as _zf
    try:
        from lxml import etree
    except ImportError:
        return {}

    result = {}
    try:
        with _zf.ZipFile(docx_path, 'r') as z:
            # 1. 解析 relationships: rId -> media 文件
            rels_xml = z.read('word/_rels/document.xml.rels')
            rels_root = etree.fromstring(rels_xml)
            ns_rel = 'http://schemas.openxmlformats.org/package/2006/relationships'
            rId_to_target = {}
            for rel in rels_root.findall(f'{{{ns_rel}}}Relationship'):
                rid = rel.get('Id', '')
                target = rel.get('Target', '')
                if 'media' in target.lower() or 'image' in target.lower():
                    rId_to_target[rid] = os.path.basename(target)

            # 2. 解析 document.xml: 遍历 body 子元素（段落+表格交替）
            doc_xml = z.read('word/document.xml')
            doc_root = etree.fromstring(doc_xml)

            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            r_ns = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

            body = doc_root.find(f'{{{w_ns}}}body')
            if body is None:
                return {}

            order_counter = 0
            table_counter = 0  # 对应 doc.tables 的索引
            last_table_idx = -1  # 最近经过的表格索引

            for elem in body:
                tag = elem.tag

                if tag == f'{{{w_ns}}}tbl':
                    # 表格：检查表格内的图片
                    for blip in elem.findall(f'.//{{{a_ns}}}blip'):
                        embed = blip.get(f'{{{r_ns}}}embed', '')
                        if embed and embed in rId_to_target:
                            fname = rId_to_target[embed]
                            if fname not in result:
                                result[fname] = (order_counter, table_counter)
                                order_counter += 1
                    table_counter += 1
                    last_table_idx = table_counter - 1

                elif tag == f'{{{w_ns}}}p':
                    # 段落：检查段落内的图片
                    for blip in elem.findall(f'.//{{{a_ns}}}blip'):
                        embed = blip.get(f'{{{r_ns}}}embed', '')
                        if embed and embed in rId_to_target:
                            fname = rId_to_target[embed]
                            if fname not in result:
                                # 图片在段落中：关联到最近的上一个表格
                                result[fname] = (order_counter, last_table_idx)
                                order_counter += 1

    except Exception:
        return {}

    return result
