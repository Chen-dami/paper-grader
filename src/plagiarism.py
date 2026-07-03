"""
抄袭检测模块 — 五维交叉验证 + 级别分类。
- 元数据: Application ID / TotalTime / revision / creator / 时间线
- 文本: 各题提示词余弦相似度
- 图片: MD5 哈希
- Excel: 数据指纹
- 综合: 加权评分 + 级别分类 + 自动判零

级别：
  < 100  → ✅ 正常，忽略
  100~200 → ⚠️ 可疑
  200~300 → 🔴 高度可疑
  ≥ 300  → 🚫 确认抄袭，自动判零
"""
import os, json, hashlib, re
from datetime import datetime
from collections import defaultdict
from difflib import SequenceMatcher
from zipfile import ZipFile
from xml.etree import ElementTree as ET

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


# ---- 级别定义 ----
PLAG_LEVELS = {
    "normal":    {"min": 0,   "max": 99,  "label": "✅ 正常",    "color": "92D050", "action": "忽略"},
    "suspicious":{"min": 100, "max": 199, "label": "⚠️ 可疑",    "color": "FFEB9C", "action": "人工复核"},
    "high":      {"min": 200, "max": 299, "label": "🔴 高度可疑","color": "FFC000", "action": "重点审查"},
    "confirmed": {"min": 300, "max": 9999,"label": "🚫 确认抄袭","color": "FF6B6B", "action": "自动判零"},
}


def classify_level(score: int) -> dict:
    """根据可疑度分数返回级别信息"""
    for key, info in PLAG_LEVELS.items():
        if info["min"] <= score <= info["max"]:
            return {"key": key, **info}
    return {"key": "normal", **PLAG_LEVELS["normal"]}


def check_all(papers_dir: str, grading_results: list = None, output_path: str = "output/查重报告.xlsx"):
    """主入口：对试卷目录下所有 docx 做全面查重。
    返回 (pairs, auto_zero_students)：
      - pairs: 所有可疑对列表（含级别信息）
      - auto_zero_students: 需自动判零的学生文件名集合（score >= 300）
    """
    docx_files = _find_docx(papers_dir)
    if len(docx_files) < 2:
        print(f"  查重需要至少2份试卷，当前 {len(docx_files)} 份")
        return [], set()

    print(f"\n  查重：{len(docx_files)} 份试卷...")

    # 1. 提取元数据 + 文件MD5
    meta_list = [_extract_meta(f) for f in docx_files]

    # 1.5 合并评分结果到元数据
    if grading_results:
        _merge_scores(meta_list, grading_results)

    # 2. 提取所有文本（正文 + 表格）
    text_map = _extract_texts(docx_files)

    # 3. 提取图片哈希
    img_map = _extract_image_hashes(docx_files)

    # 4. 提取 Excel 指纹
    excel_map = _extract_excel_fingerprints(docx_files)

    # 5. 两两比对
    pairs = []
    n = len(docx_files)
    for i in range(n):
        for j in range(i + 1, n):
            score, reasons, flags = _compare_pair(
                meta_list[i], meta_list[j],
                text_map.get(docx_files[i], {}), text_map.get(docx_files[j], {}),
                img_map.get(docx_files[i], []), img_map.get(docx_files[j], []),
                excel_map.get(docx_files[i], ""), excel_map.get(docx_files[j], ""),
            )
            # 有分数或有标记信号才记录
            if score > 0 or flags:
                pairs.append({
                    "file_a": os.path.basename(docx_files[i]),
                    "file_b": os.path.basename(docx_files[j]),
                    "student_a": meta_list[i].get("creator", "?"),
                    "student_b": meta_list[j].get("creator", "?"),
                    "score": score,
                    "reasons": reasons,
                    "flags": flags,
                })

    pairs.sort(key=lambda x: x["score"], reverse=True)

    # 6. 级别分类 + 收集判零记录
    auto_zero_students = set()
    zero_records = []  # [{student, file, score_before, reasons}]
    for p in pairs:
        level = classify_level(p["score"])
        p["level"] = level["key"]
        p["level_label"] = level["label"]
        p["level_action"] = level["action"]
        if level["key"] == "confirmed":
            auto_zero_students.add(p["file_a"])
            auto_zero_students.add(p["file_b"])
            p["auto_zero"] = True
            # 记录判零原因
            for fname in [p["file_a"], p["file_b"]]:
                m = _find_meta(meta_list, fname)
                if m:
                    zero_records.append({
                        "student": _student_label(m),
                        "file": fname,
                        "score_before": m.get("total_score", "?"),
                        "reason": "、".join(p["reasons"][:3]),
                        "peer": _student_label(
                            _find_meta(meta_list, p["file_b"] if fname == p["file_a"] else p["file_a"])
                        ),
                    })
        else:
            p["auto_zero"] = False

    # 7. 生成报告
    _generate_report(pairs, meta_list, zero_records, output_path)

    return pairs, auto_zero_students


# ============================================================
#  元数据提取
# ============================================================

def _merge_scores(meta_list: list, grading_results: list):
    """将评分结果中的总分合并到元数据中，按文件名匹配"""
    # 构建 学号/文件名 → 总分 的映射
    score_map = {}
    for r in grading_results:
        fname = r.get("_source_file", "") or os.path.basename(r.get("文件名", ""))
        sid = str(r.get("student_id", ""))
        score_map[fname] = {
            "total_score": r.get("total_score", r.get("总分", 0)),
            "student_id": sid,
            "student_name": r.get("student_name", r.get("姓名", "")),
        }
    for m in meta_list:
        f = m.get("file", "")
        info = score_map.get(f, {})
        if info:
            m["total_score"] = info["total_score"]
            if not m["student_id"]:
                m["student_id"] = info["student_id"]
            if not m["student_name"]:
                m["student_name"] = info["student_name"]


def _find_docx(directory: str) -> list:
    """递归找所有 docx"""
    files = []
    for root, _, names in os.walk(directory):
        for n in names:
            if n.endswith('.docx') and not n.startswith('~$'):
                files.append(os.path.join(root, n))
    return sorted(files)


def _extract_meta(docx_path: str) -> dict:
    """提取 docx 内部元数据 + 文件MD5"""
    # 文件整体 MD5（最可靠的抄袭判定：完全相同的文件）
    try:
        with open(docx_path, 'rb') as f:
            file_md5 = hashlib.md5(f.read()).hexdigest()
    except Exception:
        file_md5 = ""

    meta = {"file": os.path.basename(docx_path), "path": docx_path,
            "file_md5": file_md5,
            "created": "", "modified": "", "creator": "", "lastModifiedBy": "",
            "revision": 0, "total_time": 0, "application": "", "company": ""}

    try:
        with ZipFile(docx_path, 'r') as z:
            if 'docProps/core.xml' in z.namelist():
                root = ET.fromstring(z.read('docProps/core.xml'))
                ns = {'dc': 'http://purl.org/dc/elements/1.1/',
                      'dcterms': 'http://purl.org/dc/terms/',
                      'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties'}
                meta['created'] = _text(root, './/dcterms:created', ns) or ""
                meta['modified'] = _text(root, './/dcterms:modified', ns) or ""
                meta['creator'] = _text(root, './/dc:creator', ns) or ""
                meta['lastModifiedBy'] = _text(root, './/cp:lastModifiedBy', ns) or ""
                rev = _text(root, './/cp:revision', ns)
                meta['revision'] = int(rev) if rev and rev.isdigit() else 0

            if 'docProps/app.xml' in z.namelist():
                root2 = ET.fromstring(z.read('docProps/app.xml'))
                ns2 = {'': 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'}
                tt = _text(root2, './/TotalTime', ns2)
                meta['total_time'] = int(tt) if tt and tt.isdigit() else 0
                meta['application'] = _text(root2, './/Application', ns2) or ""
                meta['company'] = _text(root2, './/Company', ns2) or ""
    except Exception:
        pass

    # 学生信息从文件名推测：255102030101李杰鸿.docx
    fname = os.path.splitext(meta["file"])[0]
    sid_match = re.search(r'(\d{11,12})', fname)
    meta["student_id"] = sid_match.group(1) if sid_match else ""
    name_match = re.search(r'[一-鿿]{2,4}', fname)
    meta["student_name"] = name_match.group(0) if name_match else ""

    return meta


def _text(el, xpath, ns):
    found = el.find(xpath, ns)
    return found.text if found is not None else None


# ============================================================
#  文本提取
# ============================================================

def _extract_texts(docx_files: list) -> dict:
    """从每份 docx 提取全部文本（正文段落 + 表格内容）"""
    text_map = {}
    for f in docx_files:
        try:
            from docx import Document
            doc = Document(f)
            parts = []

            # 正文段落
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    parts.append(t)

            # 表格内容
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        t = cell.text.strip()
                        if t:
                            parts.append(t)

            full = "\n".join(parts)
            texts = {"q_all": full}
            text_map[f] = texts
        except Exception:
            text_map[f] = {}
    return text_map


# ============================================================
#  图片哈希
# ============================================================

def _extract_image_hashes(docx_files: list) -> dict:
    """提取每份 docx 中所有图片的 MD5"""
    img_map = {}
    for f in docx_files:
        hashes = []
        try:
            with ZipFile(f, 'r') as z:
                for name in z.namelist():
                    if 'media' in name and not name.endswith('/'):
                        data = z.read(name)
                        if len(data) > 5000:  # 跳过装饰小图
                            hashes.append(hashlib.md5(data).hexdigest())
        except Exception:
            pass
        img_map[f] = hashes
    return img_map


# ============================================================
#  Excel 指纹
# ============================================================

def _extract_excel_fingerprints(docx_files: list) -> dict:
    """提取嵌入 Excel 的数据指纹"""
    excel_map = {}
    for f in docx_files:
        try:
            with ZipFile(f, 'r') as z:
                for name in z.namelist():
                    if 'embedding' in name.lower() and not name.endswith('/'):
                        data = z.read(name)
                        # 找内嵌的 xlsx
                        idx = data.find(b'PK\x03\x04')
                        if idx > 4 and b'xl/' in data[idx:idx + 5000]:
                            xlsx_data = data[idx:]
                            try:
                                import io
                                df = pd.read_excel(io.BytesIO(xlsx_data))
                                # 指纹：行数 + 列名 + 前5行的值
                                fp = f"{len(df)}r_{len(df.columns)}c_"
                                fp += hashlib.md5(df.head(5).to_csv(index=False).encode()).hexdigest()[:12]
                                excel_map[f] = fp
                            except Exception:
                                pass
        except Exception:
            pass
    return excel_map


# ============================================================
#  两两比对
# ============================================================

def _compare_pair(m1: dict, m2: dict, t1: dict, t2: dict,
                  im1: list, im2: list, e1: str, e2: str) -> tuple:
    """对比一对试卷，返回 (score, reasons, flags)。
    flags 是不看分数直接标黄的可疑信号列表。"""
    score = 0
    reasons = []
    flags = []

    # --- 🔑 文件MD5完全相同（最高优先级，直接300+） ---
    if m1.get("file_md5") and m2.get("file_md5"):
        if m1["file_md5"] == m2["file_md5"]:
            score += 350
            reasons.append("🚫 文件MD5完全相同（字节级复制）")

    t1_val = m1.get("total_time", 0)
    t2_val = m2.get("total_time", 0)

    # --- 元数据 ---
    # 同一台电脑（Application ID 相同）→ 标黄
    app_same = False
    if m1.get("application") and m2.get("application"):
        if m1["application"] == m2["application"]:
            app_same = True
            score += 30
            reasons.append(f"同电脑: ID一致({m1['application'][-20:]})")

    # TotalTime 悬殊
    if max(t1_val, t2_val) > 30 and min(t1_val, t2_val) < 5:
        name_low = os.path.splitext(m1["file"])[0] if t1_val < 5 else os.path.splitext(m2["file"])[0]
        score += 15
        reasons.append(f"编辑时长悬殊: {t1_val}min vs {t2_val}min → {name_low}疑似没做")

    # 编辑时长过短 → 标黄（不看分）
    if t1_val < 5:
        flags.append(f"{os.path.splitext(m1['file'])[0]} 编辑时长仅{t1_val}min")
    if t2_val < 5:
        flags.append(f"{os.path.splitext(m2['file'])[0]} 编辑时长仅{t2_val}min")

    # revision 太低 + 时间短 → 疑似只改了名字
    if m1.get("revision", 99) <= 5 and t1_val < 10:
        score += 10
        reasons.append(f"{os.path.splitext(m1['file'])[0]} revision={m1['revision']} 疑似只改了名字")
    if m2.get("revision", 99) <= 5 and t2_val < 10:
        score += 10
        reasons.append(f"{os.path.splitext(m2['file'])[0]} revision={m2['revision']} 疑似只改了名字")

    # revision 过低 → 标黄
    if m1.get("revision", 99) <= 5:
        flags.append(f"{os.path.splitext(m1['file'])[0]} revision={m1['revision']}")
    if m2.get("revision", 99) <= 5:
        flags.append(f"{os.path.splitext(m2['file'])[0]} revision={m2['revision']}")

    # 时间线：保存时间接近 → 标黄
    time_close = False
    if m1.get("modified") and m2.get("modified"):
        try:
            t_a = datetime.fromisoformat(m1["modified"].replace('Z', '+00:00'))
            t_b = datetime.fromisoformat(m2["modified"].replace('Z', '+00:00'))
            diff_min = abs((t_b - t_a).total_seconds()) / 60
            if diff_min < 10:
                time_close = True
                score += 10
                reasons.append(f"保存时间接近: 相差{diff_min:.0f}分钟")
        except Exception:
            pass

    # --- 文本相似度 ---
    text_sim = 0
    if t1 and t2:
        all1 = t1.get("q_all", "")
        all2 = t2.get("q_all", "")
        if all1 and all2 and len(all1) > 20 and len(all2) > 20:
            sim = SequenceMatcher(None, all1, all2).ratio()
            text_sim = sim
            if sim > 0.95:
                score += 80
                reasons.append(f"全文几乎相同 {sim:.0%}")
            elif sim > 0.85:
                score += 40
                reasons.append(f"全文高度相似 {sim:.0%}")
            elif sim > 0.70:
                score += 20
                reasons.append(f"全文相似度偏高 {sim:.0%}")

    # --- 图片 MD5 ---
    shared_imgs = set(im1) & set(im2)
    if shared_imgs:
        score += 50 * len(shared_imgs)
        reasons.append(f"图片完全相同: {len(shared_imgs)}张")

    # --- Excel 指纹相同 → 标黄 ---
    excel_same = False
    if e1 and e2 and e1 == e2:
        excel_same = True
        score += 30
        reasons.append("Excel数据指纹一致")

    return score, reasons, flags


# ============================================================
#  报告生成
# ============================================================

def _generate_report(pairs: list, meta_list: list, zero_records: list, output_path: str):
    """生成查重 Excel 报告。
    Sheet 1: 可疑对（只显示有问题的）
    Sheet 2: 判零记录（谁被自动判零、为什么）
    Sheet 3: 元数据总览（每人编辑统计）"""
    wb = Workbook()

    yellow_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    hfont = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    hfill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
    border = Border(left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin"))

    level_fills = {
        "suspicious": PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid"),
        "high": PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid"),
        "confirmed": PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid"),
    }

    # ========================
    # Sheet 1: 可疑对
    # ========================
    ws = wb.active
    ws.title = "可疑对"

    headers = ["级别", "可疑度", "学生A", "总分A", "学生B", "总分B",
               "⚠️ 标记信号", "详细原因"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=col, value=h)
        c.font = hfont; c.fill = hfill; c.border = border

    for i, p in enumerate(pairs):
        row = i + 2
        score = p["score"]
        level_key = p.get("level", "normal")
        fill = level_fills.get(level_key)

        m1 = _find_meta(meta_list, p["file_a"])
        m2 = _find_meta(meta_list, p["file_b"])

        # 组装标记信号
        flag_text = "; ".join(p.get("flags", []))
        reason_text = "; ".join(p["reasons"])

        values = [
            p.get("level_label", ""),
            score,
            _student_label(m1), m1.get("total_score", "") if m1 else "",
            _student_label(m2), m2.get("total_score", "") if m2 else "",
            flag_text,
            reason_text,
        ]
        for col, v in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=v)
            c.border = border
            if fill:
                c.fill = fill
            # 标记列标黄
            if col == 7 and flag_text:
                c.fill = yellow_fill

    ws.column_dimensions['A'].width = 14
    ws.column_dimensions['B'].width = 10
    ws.column_dimensions['C'].width = 22
    ws.column_dimensions['D'].width = 8
    ws.column_dimensions['E'].width = 22
    ws.column_dimensions['F'].width = 8
    ws.column_dimensions['G'].width = 45
    ws.column_dimensions['H'].width = 60

    # ========================
    # Sheet 2: 判零记录
    # ========================
    ws2 = wb.create_sheet("判零记录")
    zheaders = ["学生", "文件名", "原得分", "判零原因", "关联人"]
    for col, h in enumerate(zheaders, 1):
        c = ws2.cell(row=1, column=col, value=h)
        c.font = hfont; c.fill = PatternFill(start_color="C00000", end_color="C00000", fill_type="solid")
        c.border = border

    zero_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    for i, z in enumerate(zero_records):
        row = i + 2
        vals = [z["student"], z["file"], z["score_before"], z["reason"], z["peer"]]
        for col, v in enumerate(vals, 1):
            c = ws2.cell(row=row, column=col, value=v)
            c.border = border
            c.fill = zero_fill

    for col in range(1, len(zheaders) + 1):
        ws2.column_dimensions[_col_letter(col)].width = 22
    ws2.column_dimensions['D'].width = 50

    if not zero_records:
        ws2.cell(row=2, column=1, value="（无）").border = border

    # ========================
    # Sheet 3: 元数据总览
    # ========================
    ws3 = wb.create_sheet("元数据")
    mheaders = ["文件名", "学号", "姓名", "总分", "文件MD5(前8位)", "Application ID(尾20)",
                "创建时间", "修改时间", "作者", "最后编辑者", "编辑时长(min)", "revision"]
    for col, h in enumerate(mheaders, 1):
        c = ws3.cell(row=1, column=col, value=h)
        c.font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
        c.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        c.border = border

    red_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    for i, m in enumerate(meta_list):
        row = i + 2
        tt = m.get("total_time", 0)
        rev = m.get("revision", 0)
        vals = [
            m.get("file", ""), m.get("student_id", ""), m.get("student_name", ""),
            m.get("total_score", ""),
            m.get("file_md5", "")[:8],
            m.get("application", "")[-20:], m.get("created", ""), m.get("modified", ""),
            m.get("creator", ""), m.get("lastModifiedBy", ""),
            tt, rev,
        ]
        for col, v in enumerate(vals, 1):
            c = ws3.cell(row=row, column=col, value=v)
            c.border = border
        # 编辑时长短 → 黄
        if isinstance(tt, (int, float)) and tt < 5:
            for col in range(1, len(mheaders) + 1):
                ws3.cell(row=row, column=col).fill = yellow_fill
        # revision低 → 红
        if isinstance(rev, (int, float)) and rev <= 5 and isinstance(tt, (int, float)) and tt < 10:
            for col in range(1, len(mheaders) + 1):
                ws3.cell(row=row, column=col).fill = red_fill

    for col in range(1, len(mheaders) + 1):
        ws3.column_dimensions[_col_letter(col)].width = 20

    d = os.path.dirname(output_path)
    if d: os.makedirs(d, exist_ok=True)
    try:
        wb.save(output_path)
    except PermissionError:
        raise PermissionError(
            f"无法保存 {output_path}，文件可能已在 Excel 中打开。请关闭后重试。"
        )
    print(f"  查重报告: {output_path}")


def _col_letter(col: int) -> str:
    """列号转字母（支持超过 Z）"""
    result = ""
    while col > 0:
        col, rem = divmod(col - 1, 26)
        result = chr(65 + rem) + result
    return result


def _find_meta(meta_list: list, filename: str) -> dict | None:
    for m in meta_list:
        if m.get("file") == filename:
            return m
    return None


def _student_label(meta: dict | None) -> str:
    if not meta:
        return "?"
    sid = meta.get("student_id", "")
    name = meta.get("student_name", "")
    return f"{sid}{name}" if sid or name else meta.get("file", "?")


def _extract_reason(reasons: list, keyword: str) -> str:
    for r in reasons:
        if keyword in r:
            return r[:40]
    return ""


# ============================================================
#  main.py 调用入口
# ============================================================

def run(papers_dir: str = "data/papers", output_dir: str = "output") -> tuple:
    """供 main.py 调用的入口。查重报告存入 output_dir 下。
    返回 (pairs, auto_zero_students)。"""
    os.makedirs(output_dir, exist_ok=True)
    report_path = os.path.join(output_dir, "查重报告.xlsx")
    return check_all(papers_dir, [], report_path)
