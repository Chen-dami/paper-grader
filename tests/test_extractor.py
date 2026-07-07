"""
extractor.py 测试 —— Word 文档解析
覆盖：段落提取、表格提取、图片提取、OLE 嵌入文件提取、学生信息识别
"""
import os
import io
import zipfile
import pytest
from unittest.mock import patch, MagicMock

from src.extractor import (
    extract, extract_from_student_folder,
    _get_image_size, _extract_ole, _try_extract_ole_native,
    _detect_office_ext, _parse_student_info,
)


class TestImageSize:
    """图片尺寸解析"""

    def test_png_size(self):
        """PNG 尺寸解析"""
        # 构造一个最小的合法 PNG（1x1 像素）
        data = (
            b'\x89PNG\r\n\x1a\n'          # PNG 签名
            b'\x00\x00\x00\rIHDR'          # IHDR chunk
            b'\x00\x00\x00\x01'            # width=1
            b'\x00\x00\x00\x01'            # height=1
            b'\x08\x02\x00\x00\x00'        # 其它
            b'\x90wS\xde'                  # CRC
        )
        w, h = _get_image_size(data, "test.png")
        assert w == 1
        assert h == 1

    def test_png_size_1920x1080(self):
        """PNG 1920x1080"""
        data = (
            b'\x89PNG\r\n\x1a\n'
            b'\x00\x00\x00\rIHDR'
            b'\x00\x00\x07\x80'            # width=1920 (0x0780)
            b'\x00\x00\x04\x38'            # height=1080 (0x0438)
            b'\x08\x02\x00\x00\x00'
            b'\x00\x00\x00\x00'            # CRC (简化)
        )
        w, h = _get_image_size(data, "test.png")
        assert w == 1920
        assert h == 1080

    def test_jpeg_returns_zero(self):
        """JPEG 返回 (0,0)（简化处理）"""
        w, h = _get_image_size(b'\xff\xd8\xff\xe0' + b'\x00' * 100, "test.jpg")
        assert w == 0
        assert h == 0

    def test_unknown_format_returns_zero(self):
        """未知格式返回 (0,0)"""
        w, h = _get_image_size(b'\x00' * 100, "test.bmp")
        assert w == 0
        assert h == 0

    def test_corrupt_data_returns_zero(self):
        """损坏数据不抛异常，返回 (0,0)"""
        w, h = _get_image_size(b'\x89PNG\x00\x00\x00', "test.png")
        assert w == 0
        assert h == 0


class TestOfficeExtDetection:
    """Office 文件类型检测"""

    def test_detect_xlsx(self):
        """检测 xlsx（含 xl/）"""
        # 最小 ZIP 结构
        data = b'PK\x03\x04' + b'\x00' * 22 + b'xl/workbook.xml\x00'
        ext = _detect_office_ext(data)
        assert ext == '.xlsx'

    def test_detect_docx(self):
        """检测 docx（含 word/）"""
        data = b'PK\x03\x04' + b'\x00' * 22 + b'word/document.xml\x00'
        ext = _detect_office_ext(data)
        assert ext == '.docx'

    def test_detect_pptx(self):
        """检测 pptx（含 ppt/）"""
        data = b'PK\x03\x04' + b'\x00' * 22 + b'ppt/presentation.xml\x00'
        ext = _detect_office_ext(data)
        assert ext == '.pptx'

    def test_detect_fallback_zip(self):
        """兜底返回 .zip"""
        ext = _detect_office_ext(b'PK\x03\x04' + b'\x00' * 100)
        assert ext == '.zip'


class TestOleNativeExtraction:
    """OLE Native 流提取"""

    def test_not_ole_returns_none(self):
        """非 OLE 数据直接返回 None"""
        result = _try_extract_ole_native(b'not an ole file')
        assert result is None

    def test_valid_ole_header_but_no_olefile(self):
        """有效的 OLE 头但无法解析（无 olefile 库时）"""
        # OLE2 签名
        data = b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1' + b'\x00' * 500
        # 没有 olefile 或解析失败，应安全返回 None
        result = _try_extract_ole_native(data)
        # 可能返回 None（解析失败）或数据（如果 olefile 可用）
        assert result is None or isinstance(result, bytes)


class TestStudentInfoParsing:
    """学生信息解析"""

    def test_from_paragraph(self):
        """从段落提取学号/姓名/班级"""
        paragraphs = [(0, "班级 软件2501 姓名 张三 学号 255102030101")]
        info = _parse_student_info(paragraphs, [])
        assert info["学号"] == "255102030101"
        assert info["姓名"] == "张三"
        assert info["班级"] == "软件2501"

    def test_from_table(self):
        """从表格提取 — 标签和值在同一行的相邻单元格"""
        tables = [[["学号", "255102030102", "姓名", "李四", "班级", "软件2501"]]]
        info = _parse_student_info([], tables)
        assert info["学号"] == "255102030102"
        assert info["姓名"] == "李四"
        assert info["班级"] == "软件2501"

    def test_from_folder_name(self):
        """从文件夹名兜底提取"""
        paragraphs = []
        docx_path = "/data/papers/软件2501/255102030103_王五/255102030103王五.docx"
        info = _parse_student_info(paragraphs, [], docx_path)
        assert info["学号"] == "255102030103"
        assert info["姓名"] == "王五"
        assert info["班级"] == "软件2501"

    def test_from_folder_no_underscore(self):
        """文件夹名无下划线"""
        docx_path = "/data/papers/电商2501/255102030101_李杰鸿/255102030101李杰鸿.docx"
        info = _parse_student_info([], [], docx_path)
        assert info["学号"] == "255102030101"
        assert info["姓名"] == "李杰鸿"

    def test_all_empty_when_nothing_found(self):
        """当所有来源都无信息时返回空"""
        info = _parse_student_info([], [], "")
        assert info["学号"] == ""
        assert info["姓名"] == ""
        assert info["班级"] == ""

    def test_priority_paragraph_over_folder(self):
        """段落信息优先于文件夹名"""
        paragraphs = [(0, "班级 测试班 姓名 优先 学号 999999999999")]
        docx_path = "/data/papers/班级A/111111111111_真实/111111111111真实.docx"
        info = _parse_student_info(paragraphs, [], docx_path)
        assert info["学号"] == "999999999999"
        assert info["姓名"] == "优先"


class TestExtractDocx:
    """extract() 主函数集成测试（需要真实 docx）"""

    def test_extract_creates_output_dirs(self, tmp_dir):
        """解析后创建正确的输出目录"""
        # 创建一个最小 docx
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_minimal_docx(docx_path, paragraphs=["学号：123  姓名：测试"])

        output_dir = os.path.join(tmp_dir, "output")
        result = extract(docx_path, output_dir)

        assert result["file_name"] == "test.docx"
        assert os.path.isdir(result["paper_dir"])
        assert os.path.isdir(os.path.join(result["paper_dir"], "images"))
        assert os.path.isdir(os.path.join(result["paper_dir"], "embeddings"))

    def test_extract_paragraphs(self, tmp_dir):
        """提取段落文本"""
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_minimal_docx(docx_path, paragraphs=[
            "第一段文字",
            "第二段：包含学号 255102030101",
        ])

        result = extract(docx_path, os.path.join(tmp_dir, "output"))
        texts = [t for _, t in result["paragraphs"] if t.strip()]
        assert "第一段文字" in texts
        assert "255102030101" in " ".join(texts)

    def test_extract_tables(self, tmp_dir):
        """提取表格"""
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_minimal_docx(docx_path, paragraphs=["测试"], table_rows=[
            ["列A", "列B", "列C"],
            ["数据1", "数据2", "数据3"],
        ])

        result = extract(docx_path, os.path.join(tmp_dir, "output"))
        assert len(result["tables"]) >= 1
        assert result["tables"][0][0] == ["列A", "列B", "列C"]

    def test_extract_images_from_docx(self, tmp_dir):
        """提取 docx 中的图片"""
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_docx_with_real_image(docx_path, tmp_dir)

        result = extract(docx_path, os.path.join(tmp_dir, "output"))
        assert len(result["images"]) >= 1
        # 每张图片返回 (path, width, height, size)
        assert len(result["images"][0]) == 4

    def test_empty_docx(self, tmp_dir):
        """空文档不崩溃"""
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_minimal_docx(docx_path, paragraphs=[])

        result = extract(docx_path, os.path.join(tmp_dir, "output"))
        assert result["file_name"] == "test.docx"
        assert result["paragraphs"] == []
        assert result["tables"] == []

    def test_extract_from_student_folder_with_supplementary(self, tmp_dir):
        """从学生文件夹提取（含辅助文件）"""
        docx_path = os.path.join(tmp_dir, "test.docx")
        _create_minimal_docx(docx_path, paragraphs=["测试内容"])

        # 创建辅助图片
        img_path = os.path.join(tmp_dir, "extra.png")
        _create_minimal_png(img_path)

        supplementary = {"images": [img_path], "excel": []}
        result = extract_from_student_folder(docx_path, os.path.join(tmp_dir, "output"), supplementary)

        assert result["file_name"] == "test.docx"
        # 辅助图片应该被注入
        assert len(result["images"]) >= 1


# ============================================================
#  辅助函数
# ============================================================

def _create_minimal_docx(path, paragraphs=None, table_rows=None):
    """创建最小 Word 文档供测试"""
    from docx import Document
    doc = Document()
    for p in (paragraphs or []):
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row_data in enumerate(table_rows):
            for j, text in enumerate(row_data):
                table.cell(i, j).text = text
    doc.save(path)


def _create_docx_with_real_image(path, tmp_dir):
    """创建一个包含真实 PNG 图片的 docx"""
    from docx import Document
    from docx.shared import Inches
    from PIL import Image as PILImage
    doc = Document()
    doc.add_paragraph("含图片的文档")

    # 创建真实的 PNG 图片
    tmp_png = os.path.join(tmp_dir, "_test_img.png")
    img = PILImage.new('RGB', (100, 100), color=(255, 0, 0))
    img.save(tmp_png, 'PNG')

    doc.add_picture(tmp_png, width=Inches(1))
    doc.save(path)


def _create_minimal_png(path):
    """创建最小 PNG 文件"""
    data = (
        b'\x89PNG\r\n\x1a\n'
        b'\x00\x00\x00\rIHDR'
        b'\x00\x00\x00\x01'
        b'\x00\x00\x00\x01'
        b'\x08\x02\x00\x00\x00'
        b'\x90wS\xde'
    )
    with open(path, 'wb') as f:
        f.write(data)
